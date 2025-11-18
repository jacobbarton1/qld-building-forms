import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
from datetime import datetime
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class InspectionFormApp:
    """
    Main application class for the inspection form application.
    Provides a GUI for filling out inspection forms with save/load functionality.
    Now generates DOCX files instead of PDFs for better compatibility with the template.
    """
    
    def __init__(self, root):
        self.root = root
        self.root.title("Inspection Form Application")
        self.root.geometry("1100x800")  # Wider window to accommodate browse buttons
        
        # Variables to track file paths
        self.current_json_path = None
        self.current_docx_path = None
        
        # Create the UI
        self.create_widgets()
        
        # Global details for building certifier and appointed competent person
        self.global_details = {"building_certifier": [], "appointed_competent_person": []}
        self.load_global_details()
        
        # Load defaults
        self.load_defaults()
        
    def create_widgets(self):
        """
        Creates all the widgets for the application interface.
        """
        # Create main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights for resizing
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Create frame for the form that will go in the main window
        form_container = ttk.Frame(main_frame)
        form_container.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        main_frame.rowconfigure(0, weight=1)
        
        # Create scrollable form frame inside the container
        canvas = tk.Canvas(form_container)
        scrollbar = ttk.Scrollbar(form_container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Grid the canvas and scrollbar in the form_container
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        # Configure the container to expand properly
        form_container.columnconfigure(0, weight=1)
        form_container.rowconfigure(0, weight=1)

        # Enable mouse wheel scrolling for different platforms
        def _on_mousewheel(event):
            # On Windows, event.delta is available; on other platforms use event.num
            if event.num == 4 or event.delta > 0:  # Scroll up
                canvas.yview_scroll(-1, "units")
            elif event.num == 5 or event.delta < 0:  # Scroll down
                canvas.yview_scroll(1, "units")

        def _bind_to_mousewheel(event):
            # Linux uses event.num, Windows/MacOS use event.delta
            canvas.bind_all("<MouseWheel>", _on_mousewheel)  # Windows
            canvas.bind_all("<Button-4>", _on_mousewheel)    # Linux
            canvas.bind_all("<Button-5>", _on_mousewheel)    # Linux

        def _unbind_from_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")

        canvas.bind('<Enter>', _bind_to_mousewheel)
        canvas.bind('<Leave>', _unbind_from_mousewheel)

        # Define form fields based on the actual template.docx structure
        self.form_field_configs = [
            # Section 1
            ("Aspect of building work (indicate the aspect)", "text"),
            
            # Section 2
            ("header2", "header", "2. Property description"),
            ("Street address", "text"),
            ("Suburb/locality", "text"),
            ("State", "text"),
            ("Postcode", "text"),
            ("Lot and plan details", "text"),
            ("Local government area the land is situated in", "text"),
            
            # Section 3
            ("header3", "header", "3. Building/structure description"),
            ("Building/structure description", "text"),
            ("Class of building/structure", "text"),
            
            # Section 4
            ("header4", "header", "4. Description of the extent of aspect/s certified"),
            ("Description of the extent of aspect/s certified", "textarea"),

            # Section 5
            ("header5", "header", "5. Basis of certification"),
            ("Basis of certification", "textarea"),

            # Section 6
            ("header6", "header", "6. Reference documentation"),
            ("Reference documentation", "textarea"),
            
            # Section 7
            ("header7", "header", "7. Building certifier reference number and building development approval number"),
            ("Building certifier's name (in full)", "text"),
            ("Building certifier reference number", "text"),
            ("Building development approval number", "text"),
            
            # Section 8
            ("header8", "header", "8. Details of appointed competent person"),
            ("Appointed competent person name (in full)", "text"),
            ("Company name (if applicable)", "text"),
            ("Contact person", "text"),
            ("Business phone number", "text"),
            ("Mobile", "text"),
            ("Email address", "text"),
            ("Postal address", "text"),
            ("Suburb/locality (postal)", "text"),
            ("State (postal)", "text"),
            ("Postcode (postal)", "text"),
            ("Licence class or registration type (if applicable)", "text"),
            ("Licence class or registration number (if applicable)", "text"),
            ("Date request to inspect received from building certifier", "text"),
            
            # Section 9
            ("header9", "header", "9. Signature of appointed competent person"),
            ("Signature Image Path", "file"),  # Changed to file type
            ("Date (signature)", "text"),
        ]

        # Create form fields based on configuration
        self.form_fields = {}
        current_row = 0  # Track the current row for grid placement
        for field_config in self.form_field_configs:
            if field_config[1] == "header":
                # Create a header label for section (field_config[2] contains the header text)
                header_label = ttk.Label(scrollable_frame, text=field_config[2], font=("TkDefaultFont", 10, "bold"))
                header_label.grid(row=current_row, column=0, columnspan=3, sticky="w", pady=(10, 5))
                current_row += 1
            else:
                # Create label and entry field
                label, field_type = field_config[0], field_config[1]
                label_widget = ttk.Label(scrollable_frame, text=f"{label}:")
                label_widget.grid(row=current_row, column=0, sticky="w", padx=(0, 10), pady=2)

                # Create entry field
                if field_type == "text":
                    entry = ttk.Entry(scrollable_frame, width=60)
                    entry.grid(row=current_row, column=1, sticky="ew", pady=2)
                    self.form_fields[label] = entry

                    # Add button to select from global details if applicable
                    if "Building certifier" in label or "competent person" in label.lower():
                        btn = ttk.Button(scrollable_frame, text="+", width=3,
                                       command=lambda l=label: self.select_global_detail(l))
                        btn.grid(row=current_row, column=2, padx=(5, 0), pady=2)

                elif field_type == "textarea":
                    # Create text widget for multiline input
                    text_widget = tk.Text(scrollable_frame, width=60, height=4, wrap=tk.WORD)
                    text_widget.grid(row=current_row, column=1, sticky="ew", pady=2)
                    self.form_fields[label] = text_widget

                    # Add scrollbar for the text widget
                    text_scrollbar = ttk.Scrollbar(scrollable_frame, orient="vertical", command=text_widget.yview)
                    text_scrollbar.grid(row=current_row, column=2, sticky="ns", padx=(5, 0), pady=2)
                    text_widget.config(yscrollcommand=text_scrollbar.set)

                elif field_type == "file":
                    # Create entry field for file path
                    entry = ttk.Entry(scrollable_frame, width=60)
                    entry.grid(row=current_row, column=1, sticky="ew", pady=2)
                    self.form_fields[label] = entry

                    # Create browse button
                    btn = ttk.Button(scrollable_frame, text="Browse", width=7,
                                   command=lambda l=label, e=entry: self.browse_file(l, e))
                    btn.grid(row=current_row, column=2, padx=(5, 0), pady=2)

                current_row += 1
                
        scrollable_frame.columnconfigure(1, weight=1)
        
        # Buttons frame
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, columnspan=3, pady=10)
        
        # Create buttons
        self.save_button = ttk.Button(button_frame, text="Save", command=self.save_form)
        self.save_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.generate_button = ttk.Button(button_frame, text="Generate DOCX", command=self.generate_docx)
        self.generate_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.load_button = ttk.Button(button_frame, text="Load", command=self.load_form)
        self.load_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.reset_button = ttk.Button(button_frame, text="Reset", command=self.reset_form)
        self.reset_button.pack(side=tk.LEFT)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E))

    def load_global_details(self):
        """
        Load global details for building certifier and appointed competent person from global.json
        """
        global_file = 'global.json'
        if os.path.exists(global_file):
            try:
                with open(global_file, 'r') as f:
                    self.global_details = json.load(f)
            except Exception as e:
                print(f"Error loading global details: {e}")
                self.global_details = {"building_certifier": [], "appointed_competent_person": []}
        else:
            # Create default global.json if it doesn't exist
            self.save_global_details()

    def save_global_details(self):
        """
        Save global details for building certifier and appointed competent person to global.json
        """
        with open('global.json', 'w') as f:
            json.dump(self.global_details, f, indent=2)
    
    def load_defaults(self):
        """
        Load default values from defaults.json file and apply to form fields
        """
        defaults_file = 'defaults.json'
        if os.path.exists(defaults_file):
            try:
                with open(defaults_file, 'r') as f:
                    defaults = json.load(f)
                
                # Apply defaults to form fields
                for field_name, default_value in defaults.items():
                    if field_name in self.form_fields:
                        widget = self.form_fields[field_name]
                        if isinstance(widget, tk.Text):
                            # For text widgets (text areas), we need to insert differently
                            widget.delete("1.0", tk.END)
                            widget.insert("1.0", default_value)
                        elif isinstance(widget, tk.Entry):
                            widget.delete(0, tk.END)
                            widget.insert(0, default_value)
                        else:
                            if hasattr(widget, 'delete') and hasattr(widget, 'insert'):
                                widget.delete(0, tk.END)
                                widget.insert(0, default_value)
                        
            except Exception as e:
                print(f"Error loading defaults: {e}")
        else:
            # Create default defaults.json if it doesn't exist
            self.create_default_defaults()
    
    def create_default_defaults(self):
        """
        Create a default defaults.json file with empty values
        """
        defaults = {
            "Aspect of building work (indicate the aspect)": "",
            "Street address": "",
            "Suburb/locality": "",
            "State": "",
            "Postcode": "",
            "Lot and plan details": "",
            "Local government area the land is situated in": "",
            "Building/structure description": "",
            "Class of building/structure": "",
            "Description of the extent of aspect/s certified": "",
            "Basis of certification": "",
            "Reference documentation": "",
            "Building certifier's name (in full)": "",
            "Building certifier reference number": "",
            "Building development approval number": "",
            "Appointed competent person name (in full)": "",
            "Company name (if applicable)": "",
            "Contact person": "",
            "Business phone number": "",
            "Mobile": "",
            "Email address": "",
            "Postal address": "",
            "Suburb/locality (postal)": "",
            "State (postal)": "",
            "Postcode (postal)": "",
            "Licence class or registration type (if applicable)": "",
            "Licence class or registration number (if applicable)": "",
            "Date request to inspect received from building certifier": datetime.now().strftime("%Y-%m-%d"),
            "Signature": "",
            "Date (signature)": datetime.now().strftime("%Y-%m-%d"),
        }
        
        with open('defaults.json', 'w') as f:
            json.dump(defaults, f, indent=2)

    def add_to_global_details(self, detail_type, detail):
        """
        Add a new detail to global details if it doesn't already exist
        """
        # Check for duplicates based on name and approval number
        is_duplicate = False
        for existing_detail in self.global_details[detail_type]:
            if (existing_detail.get("name", "") == detail.get("name", "") and
                existing_detail.get("approval_number", "") == detail.get("approval_number", "")):
                is_duplicate = True
                break
        
        if not is_duplicate:
            self.global_details[detail_type].append(detail)
            self.save_global_details()
    
    def check_and_add_to_global_details(self, form_data):
        """
        Check for new building certifier or appointed competent person details and add to global.json
        """
        # Check if there's new building certifier data
        certifier_data = {
            "name": form_data.get("Building certifier's name (in full)", ""),
            "contact": form_data.get("Building certifier reference number", ""),
            "approval_number": form_data.get("Building development approval number", "")
        }
        
        if certifier_data["name"] or certifier_data["contact"] or certifier_data["approval_number"]:
            self.add_to_global_details("building_certifier", certifier_data)
        
        # Check if there's new appointed competent person data
        person_data = {
            "name": form_data.get("Appointed competent person name (in full)", ""),
            "contact": form_data.get("Email address", ""),
            "approval_number": form_data.get("Licence class or registration number (if applicable)", "")
        }
        
        if person_data["name"] or person_data["contact"] or person_data["approval_number"]:
            self.add_to_global_details("appointed_competent_person", person_data)

    def save_form(self):
        """
        Save the current form data to a JSON file
        """
        form_data = self.get_form_data()
        
        # Determine suggested file name
        if self.current_json_path:
            initial_dir = os.path.dirname(self.current_json_path)
            initial_file = os.path.basename(self.current_json_path)
        else:
            project_name = form_data.get("Building/structure description", "inspection_form")
            if not project_name:
                project_name = form_data.get("Street address", "inspection_form")
            initial_file = f"{project_name}.json"
            initial_dir = None
        
        # Ask user for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialfile=initial_file,
            initialdir=initial_dir
        )
        
        if not file_path:
            return  # User cancelled
        
        try:
            # Save form data to JSON file
            with open(file_path, 'w') as f:
                json.dump(form_data, f, indent=2)
            
            # Update current path
            self.current_json_path = file_path
            
            # Add to global details if new building certifier or competent person data
            self.check_and_add_to_global_details(form_data)
            
            self.status_var.set(f"Form saved successfully: {file_path}")
            messagebox.showinfo("Success", f"Form saved successfully:\n{file_path}")
            
        except Exception as e:
            self.status_var.set(f"Error saving form: {str(e)}")
            messagebox.showerror("Error", f"Error saving form:\n{str(e)}")
    
    def generate_docx(self):
        """
        Generate a DOCX with the current form data
        """
        # Get form data
        form_data = self.get_form_data()

        # Determine output file path
        if self.current_json_path:
            # Suggest DOCX name based on JSON file name
            base_name = os.path.splitext(os.path.basename(self.current_json_path))[0]
            suggested_docx_path = os.path.join(os.path.dirname(self.current_json_path), f"{base_name}.docx")
        else:
            # Use project name as default
            project_name = form_data.get("Building/structure description", "inspection_form")
            if not project_name:
                project_name = form_data.get("Street address", "inspection_form")
            suggested_docx_path = f"{project_name}.docx"

        # Ask user for output file location
        output_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("DOCX files", "*.docx"), ("All files", "*.*")],
            initialfile=os.path.basename(suggested_docx_path),
            initialdir=os.path.dirname(suggested_docx_path) if os.path.dirname(suggested_docx_path) != '' else None
        )

        if not output_path:
            return  # User cancelled

        try:
            # Create a new document based on the template
            doc = Document('template.docx')

            # Extract the signature file path
            signature_path = form_data.get("Signature Image Path", "").strip()

            # Add signature image if provided
            if signature_path and os.path.exists(signature_path):
                # Add a paragraph break before signature section
                doc.add_paragraph()

                # Add a heading for signature section
                doc.add_paragraph("Signature Section:", style='Heading 1')

                # Add the signature image
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()

                # Insert the signature image, scaled appropriately
                from docx.shared import Inches
                try:
                    run.add_picture(signature_path, width=Inches(2))  # Scale to 2 inches width
                except Exception as img_error:
                    # If image can't be added, add an error message
                    doc.add_paragraph(f"Could not add signature image: {img_error}")

            # Add the form data as text
            doc.add_paragraph()
            doc.add_paragraph("FILLED FORM DATA:", style='Heading 1')

            # Add all form data except the signature path (since we handled it separately)
            for field_name, value in form_data.items():
                if field_name != "Signature Image Path" and value.strip():  # Only add non-empty values except signature path
                    doc.add_paragraph(f"{field_name}: {value}")

            # Save the document
            doc.save(output_path)

            # Create corresponding JSON file
            json_output_path = os.path.splitext(output_path)[0] + ".json"
            with open(json_output_path, 'w') as f:
                json.dump(form_data, f, indent=2)

            # Update current paths
            self.current_docx_path = output_path
            self.current_json_path = json_output_path

            self.status_var.set(f"DOCX generated successfully: {output_path}")
            messagebox.showinfo("Success", f"DOCX generated successfully:\n{output_path}")

        except Exception as e:
            self.status_var.set(f"Error generating DOCX: {str(e)}")
            messagebox.showerror("Error", f"Error generating DOCX:\n{str(e)}")

    def load_form(self):
        """
        Load form data from a JSON file
        """
        file_path = filedialog.askopenfilename(
            title="Open Form Data",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if not file_path:
            return  # User cancelled
        
        try:
            # Load form data from JSON file
            with open(file_path, 'r') as f:
                form_data = json.load(f)
            
            # Populate form fields with loaded data
            for field_name, value in form_data.items():
                if field_name in self.form_fields:
                    widget = self.form_fields[field_name]
                    if isinstance(widget, tk.Text):
                        # For text widgets (text areas), we need to insert differently
                        widget.delete("1.0", tk.END)
                        widget.insert("1.0", value)
                    elif isinstance(widget, tk.Entry):
                        widget.delete(0, tk.END)
                        widget.insert(0, value)
                    else:
                        if hasattr(widget, 'delete') and hasattr(widget, 'insert'):
                            widget.delete(0, tk.END)
                            widget.insert(0, value)
            
            # Update current path
            self.current_json_path = file_path
            
            self.status_var.set(f"Form loaded successfully: {file_path}")
            messagebox.showinfo("Success", f"Form loaded successfully:\n{file_path}")
            
            # Add to global details if new building certifier or competent person data
            self.check_and_add_to_global_details(form_data)
            
        except Exception as e:
            self.status_var.set(f"Error loading form: {str(e)}")
            messagebox.showerror("Error", f"Error loading form:\n{str(e)}")
    
    def get_form_data(self):
        """
        Get current form data as a dictionary
        """
        form_data = {}
        for field_name, widget in self.form_fields.items():
            if isinstance(widget, tk.Text):
                # For text widgets (text areas), we need to get content differently
                value = widget.get("1.0", tk.END).strip()  # Get from start to end, then strip trailing newline
            elif isinstance(widget, tk.Entry):
                value = widget.get()
            else:
                value = widget.get() if hasattr(widget, 'get') else str(widget)
            form_data[field_name] = value
        return form_data

    def reset_form(self):
        """
        Reset all form fields to empty or default values
        """
        for field_name, widget in self.form_fields.items():
            if isinstance(widget, tk.Text):
                # For text widgets (text areas), we need to delete differently
                widget.delete("1.0", tk.END)
            elif isinstance(widget, tk.Entry):
                widget.delete(0, tk.END)
            else:
                if hasattr(widget, 'delete'):
                    widget.delete(0, tk.END)

        # Reload defaults
        self.load_defaults()

        self.status_var.set("Form reset to defaults")
    
    def browse_file(self, field_name, entry_widget):
        """
        Open a file dialog to browse for a signature image
        """
        file_path = filedialog.askopenfilename(
            title="Select Signature Image",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"),
                ("PNG files", "*.png"),
                ("JPEG files", "*.jpg *.jpeg"),
                ("All files", "*.*")
            ]
        )

        if file_path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, file_path)

    def select_global_detail(self, field_name):
        """
        Open a dialog to select from global details
        """
        # Determine which type of detail to select based on field name
        if "Building certifier" in field_name:
            detail_type = "building_certifier"
        elif "competent person" in field_name.lower():
            detail_type = "appointed_competent_person"
        else:
            return  # Not a valid field for global details

        # Create selection dialog
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Select {detail_type.replace('_', ' ').title()}")
        dialog.geometry("500x300")
        dialog.transient(self.root)
        dialog.grab_set()

        # Create listbox with available details
        listbox = tk.Listbox(dialog)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Add available details to listbox
        for detail in self.global_details[detail_type]:
            listbox.insert(tk.END, detail.get("name", "") + " - " + detail.get("contact", ""))

        # Add buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=(0, 10))

        def select_detail():
            selection = listbox.curselection()
            if selection:
                selected_detail = self.global_details[detail_type][selection[0]]
                # Fill related fields based on field_name
                if "name" in field_name.lower():
                    self.form_fields[field_name].delete(0, tk.END)
                    self.form_fields[field_name].insert(0, selected_detail.get("name", ""))
                elif "reference number" in field_name.lower() or "approval number" in field_name.lower():
                    self.form_fields[field_name].delete(0, tk.END)
                    self.form_fields[field_name].insert(0, selected_detail.get("contact", ""))
                elif "contact" in field_name.lower() or "email" in field_name.lower():
                    self.form_fields[field_name].delete(0, tk.END)
                    self.form_fields[field_name].insert(0, selected_detail.get("contact", ""))
                dialog.destroy()

        def add_new():
            # For now, just close the dialog to allow manual entry
            dialog.destroy()

        ttk.Button(button_frame, text="Select", command=select_detail).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT)

        # Center the dialog
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (dialog.winfo_width() // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")

def main():
    root = tk.Tk()
    app = InspectionFormApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()